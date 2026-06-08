"""
multimodal_processor.py
───────────────────────
Unified processor for all supported file types:
  PDF, TXT  → text extraction
  JPG / PNG → OCR  (EasyOCR → pytesseract fallback)
  MP3 / WAV → speech-to-text  (OpenAI Whisper)
  MP4 / AVI → audio extraction + speech-to-text
"""
from __future__ import annotations
import os
import re
import warnings
import logging
import tempfile
from pathlib import Path
from typing import Tuple

# Must be set before transformers is imported — suppresses [transformers] __path__ warning
os.environ.setdefault("TRANSFORMERS_VERBOSITY", "error")

# Suppress noisy third-party warnings
warnings.filterwarnings("ignore", message=".*__path__.*")
warnings.filterwarnings("ignore", message=".*Accessing `__path__`.*")
logging.getLogger("transformers").setLevel(logging.ERROR)
logging.getLogger("easyocr").setLevel(logging.ERROR)
logging.getLogger("faster_whisper").setLevel(logging.WARNING)

MODALITY_MAP: dict[str, str] = {
    ".pdf":  "pdf",
    ".txt":  "text",
    ".docx": "text", ".doc": "text",  # Word documents
    ".jpg":  "image", ".jpeg": "image", ".png": "image",
    ".bmp":  "image", ".tiff": "image", ".webp": "image",
    ".mp3":  "audio", ".wav":  "audio", ".m4a":  "audio",
    ".flac": "audio", ".ogg":  "audio", ".opus": "audio",
    ".mp4":  "video", ".avi":  "video", ".mkv":  "video",
    ".mov":  "video", ".webm": "video",
}

MODALITY_ICON: dict[str, str] = {
    "pdf":   "📄",
    "text":  "📝",
    "image": "🖼️",
    "audio": "🎵",
    "video": "🎬",
}

# Model used for each modality — displayed in the UI
MODALITY_MODEL: dict[str, str] = {
    "pdf":   "PyMuPDF extract → llama3.2:3b (Q&A)",
    "text":  "Direct read → llama3.2:3b (Q&A)",
    "image": "llama3.2-vision (OCR) → RapidOCR fallback",
    "audio": "faster-whisper base (STT)",
    "video": "faster-whisper base (STT) · RapidOCR frames (>30s only)",
}


class MultimodalProcessor:
    """Extract text from any supported file type and split into chunks."""

    # Class-level model cache — loaded once, reused across all uploads
    _whisper_cache: dict = {}
    _rapidocr_cache: dict = {}

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    # ── Public API ─────────────────────────────────────────────────────────────

    def process(self, file_path: str) -> Tuple[list[str], str, str]:
        """
        Returns (chunks, full_text, modality).
        Routes each file extension to its dedicated model/extractor.

        Extension → Model mapping:
          .pdf / .txt / .docx → llama3.2:3b (Q&A on extracted text)
          .jpg / .png / …     → llama3.2-vision (OCR) → RapidOCR fallback
          .mp3 / .wav / …     → faster-whisper small (STT)
          .mp4 / .avi / …     → faster-whisper small (STT) + llama3.2-vision (frames)
        """
        path = Path(file_path)
        suffix = path.suffix.lower()
        modality = MODALITY_MAP.get(suffix)
        if modality is None:
            raise ValueError(f"Unsupported file type: {suffix!r}")

        # Route by exact suffix so formats sharing a modality can use different extractors
        suffix_extractors = {
            ".pdf":  self._extract_pdf,
            ".txt":  self._extract_txt,
            ".docx": self._extract_docx,
            ".doc":  self._extract_docx,
            ".jpg":  self._extract_image,
            ".jpeg": self._extract_image,
            ".png":  self._extract_image,
            ".bmp":  self._extract_image,
            ".tiff": self._extract_image,
            ".webp": self._extract_image,
            ".mp3":  self._extract_audio,
            ".wav":  self._extract_audio,
            ".m4a":  self._extract_audio,
            ".flac": self._extract_audio,
            ".ogg":  self._extract_audio,
            ".opus": self._extract_audio,
            ".mp4":  self._extract_video,
            ".avi":  self._extract_video,
            ".mkv":  self._extract_video,
            ".mov":  self._extract_video,
            ".webm": self._extract_video,
        }
        extractor = suffix_extractors.get(suffix)
        if extractor is None:
            raise ValueError(f"No extractor registered for: {suffix!r}")

        text = extractor(path)
        text = self._clean(text)
        chunks = self._chunk(text)
        return chunks, text, modality

    # ── Extractors ─────────────────────────────────────────────────────────────

    def _extract_pdf(self, path: Path) -> str:
        # 1. PyMuPDF — fastest, best layout recovery
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(str(path))
            pages = [page.get_text("text") for page in doc]
            doc.close()
            result = "\n\n".join(p for p in pages if p.strip())
            if result.strip():
                return result
        except ImportError:
            pass

        # 2. pdfplumber — good table/column handling
        try:
            import pdfplumber
            pages: list[str] = []
            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        pages.append(t)
            result = "\n\n".join(pages)
            if result.strip():
                return result
        except ImportError:
            pass

        # 3. PyPDF2 — legacy fallback
        import PyPDF2
        pages = []
        with open(path, "rb") as fh:
            reader = PyPDF2.PdfReader(fh)
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    pages.append(t)
        return "\n\n".join(pages)

    def _extract_txt(self, path: Path) -> str:
        return path.read_text(encoding="utf-8", errors="replace")

    def _extract_docx(self, path: Path) -> str:
        """Extract text from .docx / .doc files using python-docx."""
        try:
            from docx import Document
            doc = Document(str(path))
            parts: list[str] = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text.strip())
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        parts.append(row_text)
            return "\n\n".join(parts)
        except ImportError:
            raise ImportError(
                "python-docx is required for .docx files.\n"
                "Install it:  pip install python-docx"
            )

    def _extract_image(self, path: Path) -> str:
        """
        Image-to-text pipeline:
          1. llama3.2-vision via Ollama  (best quality, context-aware)
          2. RapidOCR via ONNX Runtime   (fast, no PyTorch)
          3. RapidOCR on a preprocessed screenshot
          4. EasyOCR                     (fallback)
          5. pytesseract                 (requires Tesseract binary)
        """
        errors: list[str] = []

        # ── 1. llama3.2-vision (Ollama) ──────────────────────────────────────────
        text = self._ollama_vision(path, errors)
        if text:
            return text

        # ── 2. RapidOCR (ONNX Runtime — no PyTorch, no system install) ──────────
        text = self._rapidocr_image(path, errors, label="RapidOCR")
        if text:
            return text

        preprocessed_paths = self._preprocess_image_for_ocr(path, errors)
        for idx, preprocessed in enumerate(preprocessed_paths, start=1):
            text = self._rapidocr_image(
                preprocessed, errors, label=f"RapidOCR preprocessed #{idx}"
            )
            if text:
                for tmp_path in preprocessed_paths:
                    tmp_path.unlink(missing_ok=True)
                return text

        # ── 3. EasyOCR ───────────────────────────────────────────────────────────
        try:
            import easyocr
            reader = easyocr.Reader(["en"], gpu=False, verbose=False)
            for ocr_path in [path, *preprocessed_paths]:
                lines = reader.readtext(str(ocr_path), detail=0, paragraph=True)
                text = " ".join(lines).strip()
                if text:
                    for tmp_path in preprocessed_paths:
                        tmp_path.unlink(missing_ok=True)
                    return text
            errors.append("EasyOCR ran but found no readable text")
        except ImportError:
            errors.append("easyocr not installed  →  pip install easyocr")
        except Exception as e:
            errors.append(f"EasyOCR failed: {e}")

        # ── 4. pytesseract (requires Tesseract binary) ───────────────────────────
        try:
            import pytesseract
            from PIL import Image as PILImage
            for ocr_path in [path, *preprocessed_paths]:
                img = PILImage.open(str(ocr_path))
                text = pytesseract.image_to_string(img).strip()
                if text:
                    for tmp_path in preprocessed_paths:
                        tmp_path.unlink(missing_ok=True)
                    return text
        except ImportError:
            errors.append("pytesseract not installed  →  pip install pytesseract Pillow")
        except Exception as e:
            errors.append(f"pytesseract failed: {e}")

        for preprocessed in preprocessed_paths:
            preprocessed.unlink(missing_ok=True)

        raise RuntimeError(
            "All image extraction methods failed:\n"
            + "\n".join(f"  • {e}" for e in errors)
            + "\n\nRecommended fix:\n"
            "  1. Start Ollama and install the vision model:\n"
            "     ollama serve\n"
            "     ollama pull llama3.2-vision\n"
            "  2. Or install the local OCR dependencies:\n"
            "     pip install rapidocr-onnxruntime Pillow easyocr"
        )

    def _rapidocr_image(self, path: Path, errors: list[str], label: str = "RapidOCR") -> str:
        """Run RapidOCR with a cached model instance and return joined text."""
        try:
            if "rapidocr" not in self.__class__._rapidocr_cache:
                from rapidocr_onnxruntime import RapidOCR
                self.__class__._rapidocr_cache["rapidocr"] = RapidOCR()
            ocr = self.__class__._rapidocr_cache["rapidocr"]
            result, _ = ocr(str(path))
            if result:
                lines = [
                    str(line[1]).strip()
                    for line in result
                    if len(line) >= 2 and str(line[1]).strip()
                ]
                text = " ".join(lines).strip()
                if text:
                    return text
            errors.append(f"{label} ran but found no readable text")
        except ImportError:
            errors.append("rapidocr-onnxruntime not installed  →  pip install rapidocr-onnxruntime")
        except Exception as e:
            errors.append(f"{label} failed: {e}")
        return ""

    def _preprocess_image_for_ocr(self, path: Path, errors: list[str]) -> list[Path]:
        """Create OCR-friendly variants for screenshots and low-contrast images."""
        try:
            from PIL import Image, ImageOps, ImageFilter
            img = Image.open(str(path)).convert("L")

            # UI screenshots often have small antialiased text. Upscale before
            # thresholding so OCR detector models can see the text regions.
            longest_side = max(img.size)
            scale = 1
            if longest_side < 1800:
                scale = 3
            elif longest_side < 3600:
                scale = 2

            if scale > 1:
                img = img.resize(
                    (img.width * scale, img.height * scale),
                    Image.Resampling.LANCZOS,
                )

            autocontrast = ImageOps.autocontrast(img).filter(ImageFilter.SHARPEN)
            threshold = autocontrast.point(lambda px: 255 if px > 170 else 0)
            inverted = ImageOps.invert(autocontrast)
            inverted_threshold = inverted.point(lambda px: 255 if px > 170 else 0)

            out_paths: list[Path] = []
            for variant in [
                autocontrast,
                threshold,
                inverted,
                inverted_threshold,
            ]:
                out = Path(tempfile.mktemp(suffix=".png"))
                variant.save(out)
                out_paths.append(out)
            return out_paths
        except Exception as e:
            errors.append(f"Image preprocessing failed: {e}")
            return []

    VISION_MODEL = "llama3.2-vision"
    OLLAMA_URL   = "http://localhost:11434"

    def _vision_model_available(self) -> bool:
        """Return True if llama3.2-vision is pulled and Ollama is reachable."""
        import urllib.request, urllib.error, json
        try:
            with urllib.request.urlopen(
                f"{self.OLLAMA_URL}/api/tags", timeout=5
            ) as resp:
                data = json.loads(resp.read().decode())
                models = [m["name"] for m in data.get("models", [])]
                return any(self.VISION_MODEL in m for m in models)
        except Exception:
            return False

    def _ollama_vision(self, path: Path, errors: list[str]) -> str:
        """Send image to llama3.2-vision via Ollama and return extracted text."""
        import base64, json, urllib.request, urllib.error

        if not self._vision_model_available():
            errors.append(
                f"{self.VISION_MODEL} not available  "
                f"→  ollama pull {self.VISION_MODEL}"
            )
            return ""

        try:
            with open(path, "rb") as f:
                img_b64 = base64.b64encode(f.read()).decode("utf-8")

            payload = json.dumps({
                "model": self.VISION_MODEL,
                "prompt": (
                    "Analyze this image thoroughly and provide ALL of the following:\n\n"
                    "1. EXTRACTED TEXT: Copy every word, number, label, caption, or symbol "
                    "visible in the image exactly as it appears. If there is no text, write 'None'.\n\n"
                    "2. DESCRIPTION: Describe everything you see in detail — objects, people, "
                    "animals, colors, shapes, layout, background, actions, expressions, "
                    "charts/graphs (including data values), diagrams, and any other visual content.\n\n"
                    "3. KEY FACTS: List the most important facts or information conveyed by this image "
                    "that someone might ask a question about.\n\n"
                    "Be thorough — your response will be used to answer questions about this image."
                ),
                "images": [img_b64],
                "stream": False,
                # Keep the vision context small enough for typical 16 GB RAM
                # systems while still leaving room for screenshot/image answers.
                "options": {
                    "temperature": 0.1,
                    "num_ctx": 4096,
                    "num_predict": 1024,
                },
            }).encode("utf-8")

            req = urllib.request.Request(
                f"{self.OLLAMA_URL}/api/generate",
                data=payload,
                headers={"Content-Type": "application/json"},
                method="POST",
            )
            with urllib.request.urlopen(req, timeout=120) as resp:
                data = json.loads(resp.read().decode("utf-8"))
                return data.get("response", "").strip()

        except urllib.error.HTTPError as e:
            detail = e.read().decode("utf-8", errors="replace").strip()
            errors.append(f"Ollama vision failed ({e.code}): {detail or e.reason}")
        except urllib.error.URLError as e:
            errors.append(f"Ollama not reachable  →  ollama serve ({e.reason})")
        except Exception as e:
            errors.append(f"llama3.2-vision failed: {e}")
        return ""

    def _extract_audio(self, path: Path) -> str:
        return self._whisper_transcribe(path)

    def _extract_video(self, path: Path) -> str:
        parts: list[str] = []
        errors: list[str] = []
        duration = self._get_video_duration(path)

        # Step 1: Audio -> faster-whisper STT when an audio track is present.
        audio_path: Path | None = None
        try:
            audio_path = self._video_to_audio(path)
            transcript = self._whisper_transcribe(audio_path)
            if transcript.strip():
                parts.append(f"Audio transcript:\n{transcript}")
        except Exception as e:
            errors.append(f"Audio transcription unavailable: {e}")
        finally:
            if audio_path and audio_path != path and audio_path.exists():
                audio_path.unlink(missing_ok=True)

        # Step 2: Frame OCR. Use frames when audio has no transcript, and also
        # sample longer clips where on-screen text is likely to matter.
        if duration > 0 and (not parts or duration >= 30):
            frame_text = self._sample_frame_ocr(path, duration)
            if frame_text.strip():
                parts.append(f"Visible frame text:\n{frame_text}")

        if parts:
            return "\n\n".join(parts)
        if errors:
            raise RuntimeError(
                "No searchable video content was extracted.\n"
                + "\n".join(f"  - {e}" for e in errors)
            )
        return ""

    def _get_video_duration(self, video_path: Path) -> float:
        """Return video duration in seconds using ffmpeg stderr output."""
        import subprocess
        try:
            ffmpeg = self._get_ffmpeg_exe()
            res = subprocess.run(
                [ffmpeg, "-i", str(video_path)],
                capture_output=True, text=True, timeout=15,
            )
            m = re.search(r"Duration:\s*(\d+):(\d+):([\d.]+)", res.stderr)
            if m:
                h, mn, s = int(m.group(1)), int(m.group(2)), float(m.group(3))
                return h * 3600 + mn * 60 + s
        except Exception:
            pass
        return 0.0

    def _sample_frame_ocr(self, video_path: Path, duration: float) -> str:
        """
        Extract evenly-spaced frames and OCR them with RapidOCR (fast, <1s/frame).
        Number of frames scales with video length — max 5 frames.
        Never calls llama3.2-vision here (too slow for batch frame processing).
        """
        import subprocess
        ffmpeg = self._get_ffmpeg_exe()

        # Scale frame count by duration. Short screen recordings still need at
        # least one visual sample, while long clips get a few evenly spaced tries.
        n_frames = max(1, min(5, int(duration / 30) + 1))

        frame_texts: list[str] = []
        for i in range(n_frames):
            t = duration * (i + 1) / (n_frames + 1)
            frame_file = Path(tempfile.mktemp(suffix=".jpg"))  # JPEG faster than PNG
            try:
                res = subprocess.run(
                    [ffmpeg, "-y", "-ss", str(t), "-i", str(video_path),
                     "-frames:v", "1", "-q:v", "5",   # lower quality = faster write
                     "-vf", "scale=1280:-1",           # preserve UI text for OCR
                     str(frame_file)],
                    capture_output=True, timeout=15,
                )
                if res.returncode == 0 and frame_file.exists():
                    text = self._fast_frame_ocr(frame_file)
                    if text.strip():
                        frame_texts.append(text.strip())
            except Exception:
                pass
            finally:
                frame_file.unlink(missing_ok=True)

        # Deduplicate identical consecutive frames
        seen: set[str] = set()
        unique: list[str] = []
        for txt in frame_texts:
            key = txt[:80]
            if key not in seen:
                seen.add(key)
                unique.append(txt)

        return "\n\n".join(unique)

    def _fast_frame_ocr(self, frame_path: Path) -> str:
        """
        OCR a single video frame using RapidOCR only.
        Does NOT call llama3.2-vision — that would take 30-60s per frame.
        """
        errors: list[str] = []
        text = self._rapidocr_image(frame_path, errors, label="RapidOCR frame")
        if text:
            return text

        preprocessed_paths = self._preprocess_image_for_ocr(frame_path, errors)
        try:
            for idx, preprocessed in enumerate(preprocessed_paths, start=1):
                text = self._rapidocr_image(
                    preprocessed,
                    errors,
                    label=f"RapidOCR frame preprocessed #{idx}",
                )
                if text:
                    return text
        finally:
            for preprocessed in preprocessed_paths:
                preprocessed.unlink(missing_ok=True)
        return ""

    # ── Audio helpers ──────────────────────────────────────────────────────────

    def _video_to_audio(self, video_path: Path) -> Path:
        """Extract audio from video as 16kHz mono WAV — optimal for Whisper."""
        import subprocess
        ffmpeg = self._get_ffmpeg_exe()   # full absolute path — avoids [WinError 2]
        out = Path(tempfile.mktemp(suffix=".wav"))

        res = subprocess.run(
            [
                ffmpeg, "-y", "-i", str(video_path),
                "-vn",                # strip video
                "-ac", "1",           # mono
                "-ar", "16000",       # 16 kHz
                "-sample_fmt", "s16", # 16-bit PCM
                str(out),
            ],
            capture_output=True, timeout=300,
        )
        if res.returncode == 0 and out.exists() and out.stat().st_size > 0:
            return out

        # Fallback: moviepy
        try:
            try:
                from moviepy import VideoFileClip
            except ImportError:
                from moviepy.editor import VideoFileClip

            clip = VideoFileClip(str(video_path))
            if clip.audio is None:
                clip.close()
                raise ValueError("Video file has no audio track.")
            clip.audio.write_audiofile(str(out), fps=16000, logger=None)
            clip.close()
            if out.exists() and out.stat().st_size > 0:
                return out
        except ImportError:
            pass

        raise RuntimeError(
            f"Could not extract audio from video '{video_path.name}'.\n"
            f"ffmpeg path used: {ffmpeg}\n"
            f"ffmpeg error: {res.stderr.decode(errors='replace')[:500]}"
        )

    def _get_ffmpeg_exe(self) -> str:
        """Return full absolute path to ffmpeg binary (system or imageio-ffmpeg bundle)."""
        import shutil
        sys_ff = shutil.which("ffmpeg")
        if sys_ff:
            return sys_ff
        try:
            import imageio_ffmpeg
            return imageio_ffmpeg.get_ffmpeg_exe()
        except ImportError:
            raise RuntimeError(
                "ffmpeg not found. Install the bundled pip package:\n"
                "  pip install imageio-ffmpeg\n"
                "Or install system ffmpeg and add it to PATH."
            )

    def _ensure_ffmpeg(self) -> None:
        """Validate ffmpeg is available (raises if not)."""
        self._get_ffmpeg_exe()  # raises if not found

    def _preprocess_audio(self, audio_path: Path) -> Path:
        """
        Normalize audio to 16 kHz mono WAV using ffmpeg.
        Whisper works best on 16kHz mono PCM — this fixes most recognition issues.
        """
        import subprocess
        ffmpeg = self._get_ffmpeg_exe()
        out = Path(tempfile.mktemp(suffix=".wav"))
        res = subprocess.run(
            [
                ffmpeg, "-y", "-i", str(audio_path),
                "-ac", "1",           # mono
                "-ar", "16000",       # 16 kHz sample rate
                "-sample_fmt", "s16", # 16-bit PCM
                "-af", "highpass=f=80,lowpass=f=8000,volume=1.5",  # denoise + boost
                str(out),
            ],
            capture_output=True, timeout=300,
        )
        if res.returncode == 0 and out.exists() and out.stat().st_size > 0:
            return out
        # If preprocessing fails, return original unchanged
        out.unlink(missing_ok=True)
        return audio_path

    def _whisper_transcribe(self, audio_path: Path) -> str:
        self._ensure_ffmpeg()

        # Normalize audio first — biggest single fix for poor recognition
        clean_path = self._preprocess_audio(audio_path)
        try:
            result = self._do_transcribe(clean_path)
        finally:
            if clean_path != audio_path and clean_path.exists():
                clean_path.unlink(missing_ok=True)
        return result

    def _do_transcribe(self, audio_path: Path) -> str:
        # faster-whisper with model caching — loads once, reused for every file
        try:
            from faster_whisper import WhisperModel
            if "faster_whisper" not in self.__class__._whisper_cache:
                self.__class__._whisper_cache["faster_whisper"] = WhisperModel(
                    "base", device="cpu", compute_type="int8"
                    # "base" loads in ~3s and transcribes 6s audio in ~2s
                    # switch to "small" only if accuracy is insufficient
                )
            model = self.__class__._whisper_cache["faster_whisper"]
            segments, info = model.transcribe(
                str(audio_path),
                beam_size=2,           # reduced for speed (was 5); still accurate
                best_of=1,             # single candidate — much faster
                temperature=0.0,
                vad_filter=True,       # skip silence — big speedup on short clips
                vad_parameters={
                    "min_silence_duration_ms": 300,
                    "speech_pad_ms": 200,
                },
                word_timestamps=False,
                language=None,
                condition_on_previous_text=False,  # faster for short audio
                no_speech_threshold=0.6,
                log_prob_threshold=-1.0,
                compression_ratio_threshold=2.4,
            )
            texts = [seg.text.strip() for seg in segments if seg.text.strip()]
            return " ".join(texts)
        except ImportError:
            pass

        # fallback: openai-whisper (cached)
        try:
            import whisper
            if "openai_whisper" not in self.__class__._whisper_cache:
                self.__class__._whisper_cache["openai_whisper"] = whisper.load_model("base")
            model = self.__class__._whisper_cache["openai_whisper"]
            result = model.transcribe(
                str(audio_path),
                fp16=False,
                beam_size=2,
                best_of=1,
                temperature=0.0,
                language=None,
                verbose=False,
            )
            return result["text"].strip()
        except ImportError:
            raise ImportError(
                "No transcription engine found. Install one:\n"
                "  pip install faster-whisper   ← recommended\n"
                "  pip install openai-whisper"
            )
        except FileNotFoundError:
            raise RuntimeError(
                "ffmpeg still not found after PATH update.\n"
                "Run:  pip install imageio-ffmpeg  and restart Streamlit."
            )

    # ── Text helpers ───────────────────────────────────────────────────────────

    def _clean(self, text: str) -> str:
        text = re.sub(r"\r\n|\r", "\n", text)
        text = re.sub(r"[^\S\n]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _chunk(self, text: str) -> list[str]:
        if not text:
            return []
        words = text.split()
        # Short content (image OCR / short transcript) — keep as single chunk.
        # Splitting short text loses context and hurts retrieval quality.
        if len(words) <= self.chunk_size:
            return [text]
        chunks: list[str] = []
        start = 0
        while start < len(words):
            end = start + self.chunk_size
            chunk = " ".join(words[start:end]).strip()
            if chunk:
                chunks.append(chunk)
            if end >= len(words):
                break
            start = end - self.chunk_overlap
        return chunks
